#!/usr/bin/env python3
"""Generate AI574_Project_Report.docx (masters-level technical report)."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


def add_para(doc: Document, text: str, style: str | None = None) -> None:
    p = doc.add_paragraph(text, style=style)
    p.paragraph_format.space_after = Pt(6)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
    for r, row in enumerate(rows):
        cells = table.rows[r + 1].cells
        for c, val in enumerate(row):
            cells[c].text = val
    doc.add_paragraph()


def main() -> None:
    repo_root = Path(__file__).resolve().parents[1]
    out_path = repo_root / "AI574_Project_Report.docx"

    doc = Document()

    # --- Title ---
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = t.add_run("AI574: Multi-Domain Corrective RAG Assistant")
    run.bold = True
    run.font.size = Pt(18)

    st = doc.add_paragraph()
    st.alignment = WD_ALIGN_PARAGRAPH.CENTER
    st.add_run(
        "Master’s Project Report — NLP & Deep Learning (Course Objective: define an NLP problem, "
        "collect/process data, build neural-network-based models, train/evaluate)"
    )

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run("Author / Team: [Your Name(s)]\nInstitution: Penn State\nDate: April 2026\nRevision: Detailed report")

    doc.add_paragraph()

    # --- Abstract ---
    doc.add_heading("Abstract", level=1)
    add_para(
        doc,
        "This report describes AI574, a multi-domain natural language assistant that combines a supervisor-routed "
        "multi-agent architecture with a shared Corrective Retrieval-Augmented Generation (CRAG) pipeline. "
        "User queries are classified into industrial control troubleshooting, recipe/cooking assistance, or "
        "scientific literature support; each specialist retrieves from a domain-isolated vector index, grades "
        "relevance, optionally rewrites the query, generates an answer conditioned on sources, and validates "
        "grounding. The system demonstrates modern NLP practice: large-scale data ingestion, dense retrieval "
        "with a pretrained transformer embedding model, and orchestrated use of large language models (LLMs) "
        "for routing, grading, generation, and validation—without training a new foundation model from scratch. "
        "We summarize datasets, preprocessing, architecture, implementation, evaluation hooks, observed index "
        "scale from notebook runs, limitations, and concrete improvements aligned with rigorous training and evaluation.",
    )

    # --- 1. Course alignment ---
    doc.add_heading("1. Introduction and Course Alignment", level=1)
    add_para(
        doc,
        "The course emphasizes defining a data-science-oriented NLP problem and solving it with deep neural networks, "
        "including dataset collection/processing, model construction, and performance evaluation. AI574 addresses a "
        "realistic industrial problem: trustworthy, domain-aware question answering over heterogeneous corpora where a "
        "single flat RAG stack often confuses vocabulary and retrieval objectives (e.g., PLC fault codes vs. recipe "
        "substitutions vs. ArXiv abstracts). The technical approach uses neural components throughout: sentence "
        "transformer embeddings for dense retrieval, and transformer-based LLMs for classification-as-routing, "
        "relevance grading, answer synthesis, and hallucination checking.",
    )
    add_para(
        doc,
        "Scope note: the primary contribution is systems-level—routing, retrieval, correction, and validation—rather "
        "than pretraining or supervised fine-tuning a new base LLM. This is common in applied NLP: pretrained encoders "
        "and generative models are composed into a testable pipeline. Section 9 proposes optional fine-tuning and "
        "benchmarking steps that would strengthen the “train and evaluate” narrative.",
    )

    # --- 2. Problem statement ---
    doc.add_heading("2. Problem Statement", level=1)
    add_para(
        doc,
        "Given a free-form user question, the system must: (1) infer the intended domain or ask for clarification; "
        "(2) retrieve relevant evidence from the correct knowledge base; (3) produce a helpful answer that stays "
        "consistent with retrieved sources; and (4) surface diagnostics (confidence, timing, sources) suitable for "
        "debugging and safe use—especially in the industrial domain.",
    )

    # --- 3. Background ---
    doc.add_heading("3. Background and Related Work", level=1)
    add_para(
        doc,
        "Retrieval-Augmented Generation (RAG) grounds LLM outputs in external documents, reducing reliance on "
        "parametric memory. Corrective RAG extends this with explicit quality control: after retrieval, documents may "
        "be filtered or the query reformulated before generation, improving robustness when the first retrieval pass "
        "misses intent. Multi-agent and router patterns reduce interference between tasks or domains. Vector databases "
        "such as ChromaDB enable scalable approximate nearest-neighbor search over dense embeddings.",
    )

    # --- 4. Data ---
    doc.add_heading("4. Datasets, Collection, and Preprocessing", level=1)

    doc.add_heading("4.1 Industrial", level=2)
    add_para(
        doc,
        "Industrial knowledge is ingested from user-provided manuals and text under a configurable data directory "
        "(e.g., PDF, Markdown, plain text). When the industrial collection is empty, seeded starter text provides "
        "a minimal demo. Chunking uses a recursive character splitter; industrial text receives abbreviation expansion "
        "and equipment-oriented normalization to improve matching for fault codes and control terminology.",
    )

    doc.add_heading("4.2 Recipe", level=2)
    add_para(
        doc,
        "The recipe domain indexes Food.com–style structured recipe data from RAW_recipes.csv (and optionally "
        "RAW_interactions.csv for popularity signals). The pipeline deduplicates, filters malformed rows, attaches "
        "metadata (e.g., dietary tags, timing), and batches embedding/upsert operations. Incremental re-indexing "
        "can skip existing chunk IDs to shorten rebuild cycles.",
    )

    doc.add_heading("4.3 Scientific", level=2)
    add_para(
        doc,
        "Scientific coverage uses the Cornell ArXiv metadata snapshot (NDJSON), with configurable category and year "
        "filters (e.g., cs.* from 2020 onward in the notebook workflow). Abstracts and metadata are chunked and "
        "embedded; live ArXiv API usage can complement bulk indexing for misses.",
    )

    doc.add_heading("4.4 Observed index scale (representative notebook run)", level=2)
    add_para(
        doc,
        "A full indexing run recorded in Multi_Domain_Agent.ipynb reported approximate Chroma document counts: "
        "~6,150 industrial chunks, ~231,311 recipe chunks, and ~687,000 scientific chunks. Exact counts depend on "
        "filters, deduplication, and re-run policies; these figures illustrate operational scale rather than a "
        "fixed dataset statistic.",
    )

    # --- 5. Methods ---
    doc.add_heading("5. System Architecture and Methodology", level=1)
    add_para(
        doc,
        "High-level control flow: START → SupervisorAgent → {industrial | recipe | scientific | clarify | fallback} "
        "→ domain agent → shared CRAGPipeline → finalize → END. Orchestration is implemented as a LangGraph workflow "
        "over a single AgentState that carries routing decisions, intermediate artifacts, sources, timing, and flags.",
    )

    doc.add_heading("5.1 CRAG loop", level=2)
    add_para(
        doc,
        "CRAG stages: retrieve top-k from the domain collection → LLM-based document grading (batch with per-document "
        "fallback on parse errors) → query rewriting and retrieval retry up to max_rewrite_attempts → response generation "
        "with source conditioning → hallucination check / grounding validation. Failed validation can escalate so the "
        "caller can restrict or review the answer.",
    )

    doc.add_heading("5.2 Neural models", level=2)
    add_para(
        doc,
        "Embeddings default to thenlper/gte-large (1024-dimensional vectors, cosine similarity in Chroma). Generation, "
        "routing, grading, rewriting, and validation use registered LLM backends: hosted Gemini (e.g., gemini_flash by "
        "default), optional Groq Llama, and local Gemma via KerasHub—selectable per query for latency/cost experiments.",
    )

    doc.add_heading("5.3 Infrastructure (Colab / persistence)", level=2)
    add_para(
        doc,
        "The Multi_Domain_Agent notebook targets Google Colab with GPU, Google Drive for project artifacts, and a "
        "local SSD Chroma directory for SQLite reliability. Rolling tar.gz snapshots back up the vector store to Drive, "
        "with integrity checks and recovery helpers—important when indices reach hundreds of thousands of chunks.",
    )

    # --- 6. Implementation ---
    doc.add_heading("6. Implementation Highlights", level=1)
    bullets = [
        "Modular repo layout: orchestration/, agents/, rag_core/, foundation/, ingestion/, evaluation/.",
        "Supervisor: JSON-structured routing with confidence thresholding and clarify path for ambiguous input.",
        "Domain agents: industrial (safety-oriented postprocessing), recipe (substitutions, dietary metadata), scientific (ArXiv-oriented behavior).",
        "Model registry and hosted adapter with auth-error failover to a configured local model ID.",
        "Unit tests with mocks for CRAG edge cases; notebooks for end-to-end validation.",
    ]
    for b in bullets:
        doc.add_paragraph(b, style="List Bullet")

    # --- 7. Training vs inference ---
    doc.add_heading("7. Training, Inference, and Evaluation Posture", level=1)
    add_para(
        doc,
        "The embedding encoder and LLMs are used in inference mode with frozen weights from public checkpoints. "
        "“Training” in the course sense is satisfied by: (a) data preparation and index construction as a supervised-free "
        "but compute-intensive pipeline; (b) hyperparameter choices (chunk size, top-k, grading thresholds, rewrite "
        "attempts); and (c) systematic evaluation of routing and answer quality. Recommended extensions (Section 9) "
        "include retrieval metric labeling, LLM-as-judge aggregates, and optional LoRA/adapter fine-tuning on a small "
        "domain-specific corpus if institutional constraints allow.",
    )

    # --- 8. Evaluation ---
    doc.add_heading("8. Evaluation", level=1)
    add_para(
        doc,
        "evaluation/metrics.py defines an Evaluator with routing accuracy over a query list (designed for expansion to "
        "~50 queries per domain plus edge cases), LLM-judge scoring on relevance/accuracy/completeness/clarity, and "
        "verify_functional checks for non-empty responses, valid confidence, and timing schema. As of this report, "
        "the curated list is partially populated (placeholders note additional queries); formal tables of accuracy and "
        "latency should be produced once the suite is frozen and run against a fixed model configuration.",
    )
    add_para(
        doc,
        "Qualitative observations (from development and prior revision notes): the supervisor threshold trades "
        "responsiveness vs. clarification rate; batch grading reduces cost but requires robust JSON parsing fallbacks; "
        "industrial preprocessing improves recall for fault-code style queries; hallucination checks often flag "
        "confident answers with weak retrieval or wrong-document paraphrase.",
    )

    # --- 9. Improvements ---
    doc.add_heading("9. Recommended Improvements and Future Work", level=1)
    improv = [
        "Complete the routing benchmark: expand to 50 queries/domain + edge cases; report overall accuracy, per-domain accuracy, confusion matrix, and latency p50/p95.",
        "Retrieval ablations: vary top-k, chunk size, and metadata filters; measure precision@k with human or LLM labels on a fixed dev set.",
        "CRAG ablations: disable or simplify grading, rewriting, or hallucination checking; quantify quality vs. latency and cost.",
        "Model comparison: Gemini Flash vs. Pro vs. Groq Llama vs. local Gemma on the same query set; track grounding failures and escalation rate.",
        "RAG quality metrics: citation coverage, answer-source consistency scores, and optional RAGAS-style automated metrics.",
        "Robustness: sanitization before index, source allowlists, prompt-injection heuristics on retrieved chunks, stronger industrial safety disclaimers.",
        "Reproducibility: pin optional dependencies in requirements.txt (e.g., langchain-text-splitters, transformers), provide a CLI or minimal API entrypoint, and version evaluation artifacts.",
        "Observability: request IDs, structured JSON logs per node, and dashboards for routing confidence and CRAG stage timings.",
        "Fallback: integrate a real web-search tool (e.g., Tavily) with ranked snippets instead of static text.",
        "Scientific depth: optional full-text PDF ingestion with layout-aware parsing where resources permit.",
        "Optional fine-tuning: small supervised dataset for routing or cross-encoder reranking to improve retrieval beyond bi-encoders.",
    ]
    for item in improv:
        doc.add_paragraph(item, style="List Bullet")

    # --- 10. Limitations ---
    doc.add_heading("10. Limitations, Ethics, and Safety", level=1)
    add_para(
        doc,
        "Industrial answers must be treated as decision support only—always follow manufacturer documentation and site "
        "safety procedures. Scientific answers depend on abstract-level retrieval; citation-heavy claims may require "
        "full text. Hosted APIs require key management; notebooks should not commit secrets. The system assumes "
        "indexed content is benign; untrusted corpora increase prompt-injection and retrieval-poisoning risk.",
    )

    # --- 11. Conclusion ---
    doc.add_heading("11. Conclusion", level=1)
    add_para(
        doc,
        "AI574 demonstrates that explicit routing, domain-scoped retrieval, and a corrective generation loop can "
        "deliver grounded, debuggable NLP behavior across heterogeneous domains using modern deep models for "
        "embeddings and language. The strongest next steps for a master’s submission are rigorous benchmarking "
        "(routing + RAG quality), documented ablations, and reproducible packaging beyond notebook-only execution.",
    )

    # --- References ---
    doc.add_heading("References (illustrative)", level=1)
    refs = [
        "Lewis, P., et al. Retrieval-Augmented Generation for Knowledge-Intensive NLP Tasks. NeurIPS 2020.",
        "LangChain / LangGraph documentation — agent and graph orchestration patterns.",
        "ChromaDB documentation — vector persistence and collections.",
        "Sentence-Transformers / GTE — dense embedding models (thenlper/gte-large).",
        "ArXiv OAI metadata — Cornell University snapshot (Kaggle / public dump).",
        "Food.com recipe data — community recipe corpus (as used in indexing scripts).",
    ]
    for r in refs:
        doc.add_paragraph(r, style="List Bullet")

    # --- Appendix A ---
    doc.add_heading("Appendix A: Key Configuration (defaults)", level=1)
    add_table(
        doc,
        ["Setting", "Default", "Module / role"],
        [
            ["routing_confidence_threshold", "0.75", "SupervisorConfig — below → clarify"],
            ["max_rewrite_attempts", "2", "RAGConfig"],
            ["relevance_threshold", "0.7", "RAGConfig — grading cutoff"],
            ["confidence_threshold", "0.6", "RAGConfig — hallucination check"],
            ["search_top_k", "5", "VectorStoreConfig"],
            ["chunk_size", "768 tokens (target)", "ChunkingConfig"],
            ["chunk_overlap_pct", "0.15", "ChunkingConfig"],
            ["embedding_model", "thenlper/gte-large", "EmbeddingConfig"],
            ["default_model_id", "gemini_flash", "run_query default LLM"],
        ],
    )

    # --- Appendix B ---
    doc.add_heading("Appendix B: Repository Map", level=1)
    add_para(
        doc,
        "agents/ — base_agent.py, industrial_agent.py, recipe_agent.py, scientific_agent.py\n"
        "config/ — settings.py, prompts.py\n"
        "evaluation/ — metrics.py\n"
        "foundation/ — llm_wrapper.py, embedding_service.py, vector_store.py, model_registry.py\n"
        "ingestion/ — document_loader.py, chunking_pipeline.py, index_builder.py\n"
        "notebooks/ — main_demo.ipynb, Multi_Domain_Agent.ipynb\n"
        "orchestration/ — state_schema.py, supervisor.py, workflow_graph.py\n"
        "rag_core/ — crag_pipeline.py, retriever.py, document_grader.py, query_rewriter.py, "
        "response_generator.py, hallucination_checker.py\n"
        "tests/ — test_crag_pipeline.py, test_security.py",
    )

    doc.save(out_path)
    print(f"Wrote {out_path}")


if __name__ == "__main__":
    main()
